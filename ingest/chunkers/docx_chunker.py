import hashlib
import os
from typing import Iterator

import tiktoken
from docx import Document

from .base import Chunk

CHUNK_SIZE = 1000
CHUNK_OVERLAP = 200
ENCODER = tiktoken.get_encoding("cl100k_base")
HEADING_STYLES = {"Heading 1", "Heading 2", "Heading 3", "Titre 1", "Titre 2", "Titre 3"}
HEADING_PREFIX = {
    "Heading 1": "#",  "Titre 1": "#",
    "Heading 2": "##", "Titre 2": "##",
    "Heading 3": "###","Titre 3": "###",
}


def _token_count(text: str) -> int:
    return len(ENCODER.encode(text))


def _split_by_tokens(text: str, chunk_size: int, overlap: int) -> list[str]:
    tokens = ENCODER.encode(text)
    chunks = []
    start = 0
    while start < len(tokens):
        end = min(start + chunk_size, len(tokens))
        chunks.append(ENCODER.decode(tokens[start:end]))
        if end == len(tokens):
            break
        start += chunk_size - overlap
    return chunks


def _table_to_md(table) -> str:
    """Convertit un tableau Word en table Markdown (première ligne = en-tête)."""
    rows = []
    for i, row in enumerate(table.rows):
        cells = [
            cell.text.strip().replace("\n", " ").replace("|", "\\|")
            for cell in row.cells
        ]
        rows.append("| " + " | ".join(cells) + " |")
        if i == 0:
            rows.append("| " + " | ".join("---" for _ in cells) + " |")
    return "\n".join(rows)


class DOCXChunker:
    def chunk_file(self, file_path: str) -> Iterator[Chunk]:
        with open(file_path, "rb") as f:
            raw_bytes = f.read()

        file_hash = hashlib.sha256(raw_bytes).hexdigest()
        source_file = os.path.basename(file_path)

        doc = Document(file_path)

        elements: list[dict] = []
        body = doc.element.body
        for child in body:
            tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
            if tag == 'p':
                for para in doc.paragraphs:
                    if para._element is child:
                        text = para.text.strip()
                        style = para.style.name if para.style else ""
                        if text:
                            prefix = HEADING_PREFIX.get(style, "")
                            content = f"{prefix} {text}" if prefix else text
                            elements.append({
                                "type": "paragraph",
                                "content": content,
                                "is_heading": style in HEADING_STYLES,
                                "style": style,
                            })
                        break
            elif tag == 'tbl':
                for table in doc.tables:
                    if table._element is child:
                        text = _table_to_md(table)
                        if text.strip():
                            elements.append({
                                "type": "table",
                                "content": text,
                                "is_heading": False,
                                "style": "Table",
                            })
                        break

        chunk_index = 0
        chunk_buffer: list[str] = []
        buffer_tokens = 0
        current_section = ""
        current_title = ""

        def flush() -> Chunk | None:
            nonlocal chunk_index
            if not chunk_buffer:
                return None
            c = Chunk(
                content="\n\n".join(chunk_buffer),
                source_file=source_file,
                page_number=1,
                chunk_index=chunk_index,
                doc_type="docx",
                section=current_section,
                title=current_title,
                file_hash=file_hash,
            )
            chunk_index += 1
            return c

        for elem in elements:
            if elem["is_heading"]:
                if "Heading 1" in elem["style"] or "Titre 1" in elem["style"]:
                    current_title = elem["content"]
                current_section = elem["content"]

            tokens = _token_count(elem["content"])

            if tokens > CHUNK_SIZE:
                c = flush()
                if c:
                    yield c
                chunk_buffer = []
                buffer_tokens = 0
                for sub in _split_by_tokens(elem["content"], CHUNK_SIZE, CHUNK_OVERLAP):
                    yield Chunk(
                        content=sub,
                        source_file=source_file,
                        page_number=1,
                        chunk_index=chunk_index,
                        doc_type="docx",
                        section=current_section,
                        title=current_title,
                        file_hash=file_hash,
                    )
                    chunk_index += 1
                continue

            if buffer_tokens + tokens > CHUNK_SIZE and chunk_buffer:
                c = flush()
                if c:
                    yield c
                overlap_buf: list[str] = []
                ot = 0
                for p in reversed(chunk_buffer):
                    t = _token_count(p)
                    if ot + t <= CHUNK_OVERLAP:
                        overlap_buf.insert(0, p)
                        ot += t
                    else:
                        break
                chunk_buffer = overlap_buf
                buffer_tokens = ot

            chunk_buffer.append(elem["content"])
            buffer_tokens += tokens

        c = flush()
        if c:
            yield c
